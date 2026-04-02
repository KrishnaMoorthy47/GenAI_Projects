from __future__ import annotations

"""
Multi-format document extractor.

Each extractor returns a list of dicts:
    [{"text": str, "source": str, "page_num": int, "chunk_index": int}]

Supported formats:
  - PDF  : pypdf (text layer); vision fallback for sparse pages
  - Image: Groq vision (JPEG, PNG, WEBP, GIF, TIFF)
  - DOCX : python-docx paragraphs + tables
  - XLSX : openpyxl sheets → structured text rows
"""

import base64
import logging
import os
from io import BytesIO
from pathlib import Path
from typing import List

logger = logging.getLogger(__name__)

IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp", ".gif", ".tiff", ".tif"}
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx"} | IMAGE_EXTENSIONS


# ── PDF ────────────────────────────────────────────────────────────────────────

def extract_pdf(path: str) -> List[dict]:
    """
    Extract text from a PDF file.
    Uses pypdf for each page. Pages with very little text (< vision_text_threshold chars)
    are re-processed via Groq vision to capture tables / scanned content.
    """
    from pypdf import PdfReader
    from chatbot.config import get_settings
    from chatbot.adapters import llm_adapter

    settings = get_settings()
    source = os.path.basename(path)
    chunks: List[dict] = []

    reader = PdfReader(path)
    sparse_pages: List[int] = []  # 0-indexed page numbers that need vision

    for page_num, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if len(text.strip()) >= settings.vision_text_threshold:
            if text.strip():
                chunks.append({
                    "text": text.strip(),
                    "source": source,
                    "page_num": page_num,
                    "chunk_index": 0,
                })
        else:
            sparse_pages.append(page_num)

    # Vision fallback for sparse pages
    if sparse_pages:
        try:
            from pdf2image import convert_from_path
            images = convert_from_path(path)
            for page_num in sparse_pages:
                if page_num >= len(images):
                    continue
                img = images[page_num]
                buf = BytesIO()
                img.save(buf, format="PNG")
                image_bytes = buf.getvalue()
                vision_text = llm_adapter.vision_extract(image_bytes)
                if vision_text.strip():
                    chunks.append({
                        "text": vision_text.strip(),
                        "source": source,
                        "page_num": page_num,
                        "chunk_index": 0,
                    })
        except Exception as exc:
            logger.warning("Vision fallback for PDF pages %s failed: %s", sparse_pages, exc)

    logger.info("PDF '%s': extracted %d page chunks (%d via vision)", source, len(chunks), len(sparse_pages))
    return chunks


# ── Image ──────────────────────────────────────────────────────────────────────

def extract_image(path_or_bytes: str | bytes) -> List[dict]:
    """
    Extract text/description from an image via Groq vision.
    Accepts a file path (str) or raw bytes.
    """
    from chatbot.adapters import llm_adapter

    if isinstance(path_or_bytes, str):
        source = os.path.basename(path_or_bytes)
        with open(path_or_bytes, "rb") as f:
            image_bytes = f.read()
    else:
        source = "image"
        image_bytes = path_or_bytes

    vision_text = llm_adapter.vision_extract(image_bytes)
    if not vision_text.strip():
        return []

    return [{"text": vision_text.strip(), "source": source, "page_num": 0, "chunk_index": 0}]


# ── DOCX ───────────────────────────────────────────────────────────────────────

def extract_docx(path: str) -> List[dict]:
    """
    Extract paragraphs and tables from a DOCX file using python-docx.
    Tables are serialised as pipe-delimited rows.
    """
    import docx  # python-docx

    source = os.path.basename(path)
    chunks: List[dict] = []
    chunk_index = 0

    doc = docx.Document(path)

    # Paragraphs
    para_texts: List[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            para_texts.append(text)

    if para_texts:
        combined = "\n".join(para_texts)
        chunks.append({
            "text": combined,
            "source": source,
            "page_num": 0,
            "chunk_index": chunk_index,
        })
        chunk_index += 1

    # Tables
    for table_idx, table in enumerate(doc.tables):
        rows_text: List[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows_text.append(" | ".join(cells))

        if rows_text:
            table_text = f"[Table {table_idx + 1}]\n" + "\n".join(rows_text)
            chunks.append({
                "text": table_text,
                "source": source,
                "page_num": 0,
                "chunk_index": chunk_index,
            })
            chunk_index += 1

    logger.info("DOCX '%s': extracted %d chunks", source, len(chunks))
    return chunks


# ── XLSX ───────────────────────────────────────────────────────────────────────

def extract_excel(path: str) -> List[dict]:
    """
    Extract sheets from an XLSX file via openpyxl.
    Each sheet becomes a chunk: header + rows formatted as key:value pairs.
    """
    import openpyxl

    source = os.path.basename(path)
    chunks: List[dict] = []

    wb = openpyxl.load_workbook(path, data_only=True)
    for sheet_idx, sheet_name in enumerate(wb.sheetnames):
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue

        # First non-empty row is treated as header
        header = [str(cell) if cell is not None else "" for cell in rows[0]]
        lines: List[str] = [f"[Sheet: {sheet_name}]"]

        for row in rows[1:]:
            cells = [str(cell) if cell is not None else "" for cell in row]
            if not any(c.strip() for c in cells):
                continue
            # key: value pairs using header
            pairs = [f"{header[i]}: {cells[i]}" for i in range(min(len(header), len(cells))) if cells[i].strip()]
            if pairs:
                lines.append(", ".join(pairs))

        if len(lines) > 1:  # more than just the sheet header
            chunks.append({
                "text": "\n".join(lines),
                "source": source,
                "page_num": sheet_idx,
                "chunk_index": 0,
            })

    logger.info("XLSX '%s': extracted %d sheet chunks", source, len(chunks))
    return chunks


# ── Router ─────────────────────────────────────────────────────────────────────

def route(path: str, file_bytes: bytes | None = None) -> List[dict]:
    """
    Dispatch to the appropriate extractor based on file extension.
    file_bytes is only needed when path is a temp path and source name matters.
    """
    ext = Path(path).suffix.lower()

    if ext == ".pdf":
        return extract_pdf(path)
    elif ext in IMAGE_EXTENSIONS:
        return extract_image(path)
    elif ext == ".docx":
        return extract_docx(path)
    elif ext == ".xlsx":
        return extract_excel(path)
    else:
        raise ValueError(f"Unsupported file extension: '{ext}'. Supported: {SUPPORTED_EXTENSIONS}")
